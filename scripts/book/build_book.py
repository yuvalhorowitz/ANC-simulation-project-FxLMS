"""
Build the Project Book .docx from the TAU template + content/*.md files.

Strategy:
1. Copy the official TAU template into output/book/project_book.docx — preserves
   styles (Heading 1/2, Caption, table-of-figures), the TAU logo, headers/footers,
   and Word fields for the auto-TOC and auto-figure-numbering.
2. Walk the title-page tables and fill in the placeholders (project title,
   project number, student names + IDs, supervisor, project location).
3. Strip the Hebrew preamble paragraphs and the boilerplate body content from
   the template, leaving just: title page, TOC, lists.
4. Parse our 9 content .md files (abstract + 8 chapters) with a small Markdown
   subset and append paragraphs with the template's named styles. Embed PNGs
   as inline pictures with auto-numbered captions.

Run: python scripts/book/build_book.py
Output: output/book/project_book.docx (open in Word, press F9 to refresh TOC).
"""

import re
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


# Paths
ROOT = Path(__file__).resolve().parent.parent.parent
TEMPLATE = ROOT / 'progress files and requierments' / 'Book Requierments' / 'Template Project Report (eng) (1).docx'
OUT_DIR = ROOT / 'output' / 'book'
OUT_DOCX = OUT_DIR / 'project_book.docx'
CONTENT_DIR = ROOT / 'scripts' / 'book' / 'content'

# Title page metadata
META = {
    'project_title_main': 'Active Noise Reduction System In a Vehicle',
    'project_title_sub':  'FxLMS-Based ANC for Automotive Applications: '
                          'Multi-Channel Adaptive Filtering and Spatial Quiet Zones',
    'project_number':     '25-1-1-3214',
    'student_1_name':     'Ariel Turnowski',
    'student_1_id':       '206483513',
    'student_2_name':     'Yuval Horowitz',
    'student_2_id':       '206587719',
    'supervisor':         'Dr. Lior Arbel',
    'project_at':         'Tel Aviv University, School of Electrical Engineering',
}

# Chapter content files (in order) — first one becomes "Abstract" (Title style),
# the rest become Heading 1 chapters.
CHAPTERS = [
    ('abstract',                    'Abstract'),
    ('01_introduction',             'Introduction'),
    ('02_theoretical_background',   'Theoretical Background'),
    ('03_simulation',               'Simulation'),
    ('04_implementation',           'Implementation'),
    ('05_results',                  'Analysis of Results'),
    ('06_conclusions',              'Conclusions and Further Work'),
    ('07_documentation',            'Project Documentation'),
    ('08_references',               'References'),
]

# Max width for embedded pictures (inches). 4.8 keeps figures compact so
# the document fits the 20-page budget; the heatmap and scenario table are
# still legible at this size.
PICTURE_WIDTH_IN = 4.8


# ---------- Template manipulation helpers ----------

def remove_paragraph(p):
    """Remove a paragraph element from its parent."""
    p._element.getparent().remove(p._element)
    p._p = p._element = None


def replace_run_text(run, new_text):
    """Replace text content in a run while preserving its formatting."""
    run.text = new_text


def fill_title_page(doc):
    """Walk every table cell and replace placeholder strings with metadata."""
    # The template uses tables for the title-page layout. We find specific
    # placeholder strings and either replace them or fill empty cells next to
    # known labels.
    label_to_value = {
        'Insert the': META['project_title_main'],
        'Project Title': '',
        'here': '',
        'Project Number:': META['project_number'],
        'Project Report': 'Project Report',
        'Supervisor:':    META['supervisor'],
        'Project Carried Out at:': META['project_at'],
    }

    # Track which labels we've seen so we know which cells to fill
    for tbl_idx, tbl in enumerate(doc.tables):
        for row_idx, row in enumerate(tbl.rows):
            cells = row.cells
            for col_idx, cell in enumerate(cells):
                txt = cell.text.strip()
                # Replace the project-title placeholder cell entirely
                if txt.startswith('Insert the') and 'Project Title' in cell.text:
                    # Wipe paragraphs in the cell, write title + subtitle
                    for p in list(cell.paragraphs):
                        if p.text.strip():
                            for r in list(p.runs):
                                r.text = ''
                    # Use the existing first paragraph as the title carrier
                    p0 = cell.paragraphs[0]
                    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if not p0.runs:
                        r = p0.add_run(META['project_title_main'])
                    else:
                        p0.runs[0].text = META['project_title_main']
                    p0.runs[0].bold = True
                    p0.runs[0].font.size = Pt(28)
                    # Add subtitle
                    p1 = cell.add_paragraph()
                    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r1 = p1.add_run(META['project_title_sub'])
                    r1.italic = True
                    r1.font.size = Pt(13)

                # Project Number cell — the empty sibling cell is the one to fill.
                if txt == 'Project Number:':
                    # Fill the cell to the LEFT (which is the underlined blank).
                    if col_idx > 0:
                        sib = cells[col_idx - 1]
                        if not sib.text.strip():
                            sib.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            r = sib.paragraphs[0].add_run(META['project_number'])
                            r.bold = True
                            r.font.size = Pt(20)

                # Student rows: label "Student:" / "ID:" with empty sibling cells
                if txt == 'Student:':
                    # Find which student we're on by counting "Student:" occurrences
                    # in earlier rows of this table.
                    student_idx = sum(
                        1 for r in tbl.rows[:row_idx]
                        for c in r.cells if c.text.strip() == 'Student:'
                    )
                    name_key = f'student_{student_idx + 1}_name'
                    id_key = f'student_{student_idx + 1}_id'
                    if name_key in META:
                        # The cell to the right is the name slot
                        if col_idx + 1 < len(cells):
                            target = cells[col_idx + 1]
                            if not target.text.strip():
                                target.paragraphs[0].add_run(META[name_key])
                        # The next "ID:" label and its right neighbor
                        # are also in the same row.
                        for cidx in range(col_idx + 1, len(cells)):
                            if cells[cidx].text.strip() == 'ID:':
                                if cidx + 1 < len(cells):
                                    id_cell = cells[cidx + 1]
                                    if not id_cell.text.strip():
                                        id_cell.paragraphs[0].add_run(
                                            META[id_key]
                                        )
                                break

                if txt == 'Supervisor:':
                    if col_idx + 1 < len(cells):
                        target = cells[col_idx + 1]
                        if not target.text.strip():
                            target.paragraphs[0].add_run(META['supervisor'])

                if txt == 'Project Carried Out at:':
                    if col_idx + 1 < len(cells):
                        target = cells[col_idx + 1]
                        if not target.text.strip():
                            target.paragraphs[0].add_run(META['project_at'])


def strip_template_body(doc):
    """Remove the Hebrew preamble (paragraphs 0-17), the dotted auto-TOC
    field (we replace it with a simple static chapter list), the
    'List of figures / tables / equations' sections (they ship with stale
    example entries that look broken), and the example-content body
    (paragraphs from 'Abstract' Title onward).
    """
    # 1) Remove the Hebrew preamble: paragraphs 0..17 contain instructions
    #    (in Hebrew) that aren't part of the actual book.
    paras = list(doc.paragraphs)
    for p in paras[:18]:
        remove_paragraph(p)

    # 2) Remove the auto-TOC SDT block (the user reported its dot-leader
    #    alignment looked bad; we replace it with a clean static list later).
    body = doc.element.body
    for sdt in list(body.findall(qn('w:sdt'))):
        body.remove(sdt)

    # 3) Remove the "List of figures / tables / equations" stub paragraphs.
    #    They contain stale example entries from the template
    #    ("Figure 1: Block diagram\t2", "Table 1: Results\t8", ...) that we
    #    don't want shown.
    paras = list(doc.paragraphs)
    drop_keywords = ('List of figures', 'List of tables', 'List of equations')
    for p in paras:
        if (p.text.strip() in drop_keywords
                or (p.style and p.style.name == 'table of figures')):
            remove_paragraph(p)

    # 4) Remove the body content from 'Abstract' (Title style) onward.
    paras = list(doc.paragraphs)  # refresh after removals
    cut_idx = None
    for i, p in enumerate(paras):
        if p.style and p.style.name == 'Title' and p.text.strip() == 'Abstract':
            cut_idx = i
            break
    if cut_idx is not None:
        for p in paras[cut_idx:]:
            remove_paragraph(p)


def add_contents_page(doc):
    """Insert a clean, simple Contents page just before the first chapter.
    No dot leaders, no page numbers — just a numbered chapter list using
    the template's own paragraph styles. Visually minimal by design.
    """
    p = doc.add_paragraph(style='TOC Heading')
    p.add_run('Contents')

    entries = [
        ('Abstract', None),
        ('Introduction', '1'),
        ('Theoretical Background', '2'),
        ('Simulation', '3'),
        ('Implementation', '4'),
        ('Analysis of Results', '5'),
        ('Conclusions and Further Work', '6'),
        ('Project Documentation', '7'),
        ('References', '8'),
    ]
    for title, num in entries:
        line = f'{num}.  {title}' if num else title
        para = doc.add_paragraph(style='Normal')
        para.paragraph_format.left_indent = Inches(0.4)
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run(line)
        run.font.size = Pt(13)

    # Page break so the Abstract starts on a fresh page
    pb = doc.add_paragraph()
    pb.add_run().add_break()


# ---------- Markdown-lite parser ----------

IMG_RE = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')
TABLE_LINE_RE = re.compile(r'^\s*\|.*\|\s*$')
TABLE_SEP_RE = re.compile(r'^\s*\|[\s:|-]+\|\s*$')


def parse_markdown(md_text):
    """Return a flat list of block tuples we know how to render:
       ('h1', text), ('h2', text), ('h3', text),
       ('p', text),
       ('img', alt, path),
       ('table', rows: list[list[str]]),
       ('list', items: list[str]).
    """
    blocks = []
    lines = md_text.splitlines()
    i = 0
    para = []
    list_buf = []
    table_buf = []

    def flush_para():
        if para:
            blocks.append(('p', ' '.join(para).strip()))
            para.clear()

    def flush_list():
        if list_buf:
            blocks.append(('list', list_buf.copy()))
            list_buf.clear()

    def flush_table():
        if len(table_buf) >= 2:
            # Skip the separator row
            rows = []
            for j, row_line in enumerate(table_buf):
                if j == 1 and TABLE_SEP_RE.match(row_line):
                    continue
                cells = [c.strip() for c in row_line.strip().strip('|').split('|')]
                rows.append(cells)
            blocks.append(('table', rows))
        elif table_buf:
            for line in table_buf:
                para.append(line)
            flush_para()
        table_buf.clear()

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Headings
        if stripped.startswith('### '):
            flush_para(); flush_list(); flush_table()
            blocks.append(('h3', stripped[4:].strip()))
        elif stripped.startswith('## '):
            flush_para(); flush_list(); flush_table()
            blocks.append(('h2', stripped[3:].strip()))
        elif stripped.startswith('# '):
            flush_para(); flush_list(); flush_table()
            blocks.append(('h1', stripped[2:].strip()))

        # Image
        elif stripped.startswith('![') and IMG_RE.match(stripped):
            flush_para(); flush_list(); flush_table()
            m = IMG_RE.match(stripped)
            blocks.append(('img', m.group(1), m.group(2)))

        # Table row
        elif TABLE_LINE_RE.match(line):
            flush_para(); flush_list()
            table_buf.append(line)

        # List item
        elif stripped.startswith('- '):
            flush_para(); flush_table()
            list_buf.append(stripped[2:])

        # Blank line
        elif not stripped:
            flush_para(); flush_list(); flush_table()

        else:
            flush_list(); flush_table()
            para.append(stripped)

        i += 1

    flush_para(); flush_list(); flush_table()
    return blocks


# ---------- Block writers ----------

def write_paragraph(doc, text, style='Normal'):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


def write_heading(doc, text, level):
    """level: 1 for chapter, 2 for section, 3 for subsection."""
    style = f'Heading {level}'
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


def write_image(doc, alt, image_path):
    """Embed a picture and add a caption paragraph using the template's
    'Caption' style so it's automatically picked up by the auto-numbering."""
    full_path = (ROOT / image_path).resolve() if not Path(image_path).is_absolute() \
        else Path(image_path)
    if not full_path.exists():
        # Fallback: write a placeholder paragraph so the missing image is visible.
        write_paragraph(doc, f'[MISSING IMAGE: {image_path}]', style='Normal')
        return

    pic_par = doc.add_paragraph()
    pic_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_par.add_run()
    try:
        run.add_picture(str(full_path), width=Inches(PICTURE_WIDTH_IN))
    except Exception as e:
        write_paragraph(doc, f'[IMAGE INSERT FAILED: {image_path} — {e}]')
        return

    # Add caption paragraph in the template's 'Caption' style. Word's auto
    # figure numbering picks this up and refreshes the figure number on F9.
    if alt.strip():
        cap = doc.add_paragraph(style='Caption')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.add_run(alt.strip())


def write_table(doc, rows):
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=n_cols)
    tbl.style = 'Table Grid'
    for ri, row in enumerate(rows):
        for ci in range(n_cols):
            cell = tbl.cell(ri, ci)
            cell.paragraphs[0].text = ''
            txt = row[ci] if ci < len(row) else ''
            run = cell.paragraphs[0].add_run(txt)
            if ri == 0:
                run.bold = True


def write_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Paragraph')
        p.add_run(item)


# ---------- Main pipeline ----------

def append_chapter(doc, md_path, chapter_title, is_abstract=False):
    """Write a chapter heading and then render the markdown blocks."""
    if is_abstract:
        # Abstract uses the 'Title' style in the template (paragraph 44)
        p = doc.add_paragraph(style='Title')
        p.add_run(chapter_title)
    else:
        write_heading(doc, chapter_title, 1)

    if not md_path.exists():
        write_paragraph(doc, f'[CONTENT FILE MISSING: {md_path.name}]')
        return

    blocks = parse_markdown(md_path.read_text(encoding='utf-8'))
    for block in blocks:
        kind = block[0]
        if kind == 'p':
            write_paragraph(doc, block[1])
        elif kind == 'h1':
            # Within a chapter file, 'h1' is treated as level 2 (section)
            write_heading(doc, block[1], 2)
        elif kind == 'h2':
            write_heading(doc, block[1], 2)
        elif kind == 'h3':
            write_heading(doc, block[1], 3)
        elif kind == 'img':
            write_image(doc, block[1], block[2])
        elif kind == 'table':
            write_table(doc, block[1])
        elif kind == 'list':
            write_list(doc, block[1])


def main():
    if not TEMPLATE.exists():
        raise SystemExit(f'Template not found: {TEMPLATE}')

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy(TEMPLATE, OUT_DOCX)
    print(f'Copied template -> {OUT_DOCX}')

    doc = Document(str(OUT_DOCX))

    print('Filling title-page placeholders ...')
    fill_title_page(doc)

    print('Stripping template Hebrew preamble + example body ...')
    strip_template_body(doc)

    print('Adding clean Contents page ...')
    add_contents_page(doc)

    print('Appending content chapters:')
    for slug, title in CHAPTERS:
        md_path = CONTENT_DIR / f'{slug}.md'
        is_abstract = (slug == 'abstract')
        print(f'  - {title:30s} <- {md_path.name}')
        append_chapter(doc, md_path, title, is_abstract=is_abstract)

    doc.save(str(OUT_DOCX))

    size_mb = OUT_DOCX.stat().st_size / (1024 * 1024)
    print(f'\nSaved: {OUT_DOCX} ({size_mb:.2f} MB)')
    if size_mb > 10:
        print('  WARNING: file is over the 10 MB course limit; downsize images.')

    print('\nNext: open in Word, press F9 (Update Field) to refresh TOC and '
          'figure numbers, then export to PDF.')


if __name__ == '__main__':
    main()
